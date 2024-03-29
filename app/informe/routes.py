from flask import (
    render_template,
    request,
    redirect,
    url_for,
    make_response,
    flash,
    Response,
)
from werkzeug.urls import url_parse

from flask_login import login_required

from sqlalchemy import select

from . import informe_bp

from .models import Informe
from app.miembro.models import Miembro, miembros_comisiones
from app.comision.models import Comision

from .forms import CreateInformeForm

from app import db

from pdfkit import from_string

import datetime

import io
import csv

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from docx.enum.section import WD_SECTION

from config.default import BASE_DIR
import os


@informe_bp.route("/informes", methods=["GET", "POST"])
@login_required
def generate_informe():
    """
    Ruta para generar un informe
    """
    form = CreateInformeForm()

    # Las opciones del select serán todos los miembros ya estén activos o no
    # Le pasamos una tupla donde el valor será la ID del miembro y la etiqueta su nombre y apellidos
    miembros = [(m.id, m.nombre, m.apellidos) for m in Miembro.query.all()]
    miembros = [(m[0], m[1] + " " + m[2]) for m in miembros]
    form.miembro.choices = miembros
    if request.method == "POST" and form.validate_on_submit():
        secretario = form.secretario.data
        fecha_inicio = form.fecha_inicio.data
        fecha_fin = form.fecha_fin.data
        tipo_informe = form.tipo_informe.data  # 'escuela' o 'comision'
        miembro = form.miembro.data  # ID_Miembro
        tratamiento = form.tratamiento.data
        formato_informe = form.formato_informe.data

        # Creamos el informe y lo guardamos
        informe = Informe(
            secretario=secretario,
            fecha_inicio=fecha_inicio,
            fecha_fin=fecha_fin,
            tipo=tipo_informe,
            id_miembro=miembro,
        )

        informe.save()

        # Abrimos una nueva pestaña en la que el usuario verá el PDF generado para el miembro elegido
        # Comprobamos si se ha pasado por la URL el parámetro next
        next_page = request.args.get("next", None)
        if not next_page or url_parse(next_page).netloc != "":
            if tipo_informe == "comision":
                if formato_informe == "pdf":
                    next_page = url_for(
                        "informe.informe_comisiones",
                        secretario=secretario,
                        fecha_inicio=fecha_inicio,
                        fecha_fin=fecha_fin,
                        id_miembro=miembro,
                        tratamiento=tratamiento,
                    )
                elif formato_informe == "docx":
                    next_page = url_for(
                        "informe.informe_comisiones_docx",
                        secretario=secretario,
                        fecha_inicio=fecha_inicio,
                        fecha_fin=fecha_fin,
                        id_miembro=miembro,
                        tratamiento=tratamiento,
                    )
            elif tipo_informe == "escuela":
                if formato_informe == "pdf":
                    next_page = url_for(
                        "informe.informe_escuela",
                        secretario=secretario,
                        fecha_inicio=fecha_inicio,
                        fecha_fin=fecha_fin,
                        id_miembro=miembro,
                        tratamiento=tratamiento,
                    )
                elif formato_informe == "docx":
                    next_page = url_for(
                        "informe.informe_escuela_docx",
                        secretario=secretario,
                        fecha_inicio=fecha_inicio,
                        fecha_fin=fecha_fin,
                        id_miembro=miembro,
                        tratamiento=tratamiento,
                    )
        return redirect(next_page)

    return render_template("informe/generateInforme_view.html", form=form)


@informe_bp.route(
    "/informe/comisiones/<secretario>/<fecha_inicio>/<fecha_fin>/<tratamiento>/<int:id_miembro>"
)
@login_required
def informe_comisiones(secretario, fecha_inicio, fecha_fin, tratamiento, id_miembro):
    """
    Ruta que genera un informe de pertenencia a comisiones de un miembro en formato 'pdf'
    """
    # A partir de la ID del miembro recuperamos todos sus datos y las comisiones a las
    # que ha perteneceido durante todo ese tiempo
    miembro = Miembro.get_by_id(id_miembro)

    # Ahora recuperamos todas las comisiones a las que ha pertenecido durante todo ese tiempo
    comisiones = db.session.execute(
        select(
            miembros_comisiones.columns.id_comision,
            miembros_comisiones.columns.fecha_incorporacion,
            miembros_comisiones.columns.fecha_baja,
            miembros_comisiones.columns.cargo,
        ).where(
            miembros_comisiones.columns.id_miembro == id_miembro,
            miembros_comisiones.columns.fecha_incorporacion.between(
                fecha_inicio, fecha_fin
            ),
        )
    ).fetchall()

    # Si ese miembro no pertenece a ninguna comisión no generamos certificado
    if not comisiones:
        flash(
            f"El miembro {miembro.nombre} {miembro.apellidos} no ha pertenecido a ninguna comisión en el rango de fechas indicado.",
            "error",
        )
        return redirect(url_for("informe.generate_informe"))

    # Obtenemos el nombre de esas comisiones a partir de la ID
    # Primero obtenemos todas las ID de las comisiones de la lista anterior
    id_comisiones = [c[0] for c in comisiones]
    # Ahora recuperamos los nombres
    nombres_comisiones = (
        db.session.query(Comision.id, Comision.nombre)
        .filter(Comision.id.in_(id_comisiones))
        .all()
    )
    # Creamos una lista de tuplas con las comisiones, fechas, cargos y nombres utilizando map() para agregar el nombre de la comisión
    # comisiones_con_nombre -> (ID_comision, fecha_incorporacion, fecha_baja, cargo_miembro, nombre_comision)
    comisiones_con_nombre = list(
        map(
            lambda c: (
                c[0],
                c[1],
                c[2],
                c[3],
                next((n[1] for n in nombres_comisiones if n[0] == c[0]), ""),
            ),
            comisiones,
        )
    )

    # Ordenamos la lista de manera ascendente en función de la fecha de inicio
    comisiones_con_nombre = sorted(comisiones_con_nombre, key=lambda tupla: tupla[1])

    # Borramos la comisión 'Junta de Escuela' si existe ya que no pertenece a este informe
    for tupla in comisiones_con_nombre:
        if tupla[4] == "Junta de Escuela":
            comisiones_con_nombre.remove(tupla)

    # Obtenemos el HTML
    html = render_template(
        "informe/blueprintComisiones_view.html",
        secretario=secretario,
        tratamiento=tratamiento,
        miembro=miembro,
        comisiones=comisiones_con_nombre,
        fecha_actual=datetime.datetime.now(),
    )

    # PDF options
    options = {
        "page-size": "A4",
        "margin-top": "1.0cm",
        "margin-right": "1.0cm",
        "margin-bottom": "4.0cm",
        "margin-left": "1.0cm",
        "encoding": "UTF-8",
        "enable-local-file-access": "",
        # "header-html": os.path.join(
        #     BASE_DIR, "app", "informe", "templates", "informe", "headerPDF.html"
        # ),
    }

    # Construimos el PDF a partir del HTML
    pdf = from_string(html, options=options)

    # Descargamos el PDF
    # return Response(pdf, mimetype="application/pdf")
    response = make_response(pdf)
    response.headers["Content-Type"] = "application/pdf"
    response.headers["Content-Disposition"] = "inline; filename=certificado.pdf"
    return response


@informe_bp.route(
    "/informe/escuela/<secretario>/<fecha_inicio>/<fecha_fin>/<tratamiento>/<int:id_miembro>"
)
@login_required
def informe_escuela(secretario, fecha_inicio, fecha_fin, tratamiento, id_miembro):
    """
    Ruta que genera un informe de pertenencia a la junta de escuela de un miembro en formato 'pdf'
    """
    miembro = Miembro.get_by_id(id_miembro)

    # Aquí sólo nos importa la comisión "Junta de Escuela"
    # Recuperamos el ID de esta comisión
    id_comision = Comision.query.filter_by(nombre="Junta de Escuela").first().id

    if not id_comision:
        flash(
            "Parece que la comisión 'Junta de Escuela' no se encuentra registrada en el sistema.",
            "error",
        )
        return redirect(url_for("informe.generate_informe"))

    # Si este miembro está en esa comsión recuperamos sus fechas de incorporación y de baja
    # Ahora la fecha de incorporación no es necesario que esté entre los rangos que nos indica
    # Ahora nos es irrelevante las fechas que nos indique
    comision = db.session.execute(
        select(
            miembros_comisiones.columns.fecha_incorporacion,
            miembros_comisiones.columns.fecha_baja,
        ).where(
            miembros_comisiones.columns.id_miembro == id_miembro,
            miembros_comisiones.columns.id_comision == id_comision,
        )
    ).first()

    if not comision:
        flash(
            f"Parece que el miembro {miembro.nombre} {miembro.apellidos} no ha pertenecido a la Junta de la Escuela en el rango de fechas indicado.",
            "error",
        )
        return redirect(url_for("informe.generate_informe"))

    # Obtenemos el HTML
    html = render_template(
        "informe/blueprintEscuela_view.html",
        secretario=secretario,
        tratamiento=tratamiento,
        miembro=miembro,
        comision=comision,
        fecha_actual=datetime.datetime.now(),
    )

    # PDF options
    options = {
        "page-size": "A4",
        "margin-top": "1.0cm",
        "margin-right": "1.0cm",
        "margin-bottom": "4.0cm",
        "margin-left": "1.0cm",
        "encoding": "UTF-8",
        "enable-local-file-access": "",
    }

    # Construimos el PDF a partir del HTML
    pdf = from_string(html, options=options)

    # Descargamos el PDF
    # return Response(pdf, mimetype="application/pdf")
    response = make_response(pdf)
    response.headers["Content-Type"] = "application/pdf"
    response.headers["Content-Disposition"] = "inline; filename=certificado.pdf"
    return response


@informe_bp.route(
    "/informe/comisiones/docx/<secretario>/<fecha_inicio>/<fecha_fin>/<tratamiento>/<int:id_miembro>"
)
@login_required
def informe_comisiones_docx(
    secretario, fecha_inicio, fecha_fin, tratamiento, id_miembro
):
    """
    Ruta que genera un informe de pertenencia a comisiones de un miembro en formato '.docx'
    """
    miembro = Miembro.get_by_id(id_miembro)  # Aquí ya tenemos todos sus datos

    # Ahora recuperamos todas las comisiones a las que ha pertenecido durante todo ese tiempo
    # Usamos 'select' de SQLAlchemy aquí para mejorar la claridad de la setencia
    comisiones = db.session.execute(
        select(
            miembros_comisiones.columns.id_comision,
            miembros_comisiones.columns.fecha_incorporacion,
            miembros_comisiones.columns.fecha_baja,
            miembros_comisiones.columns.cargo,
        ).where(
            miembros_comisiones.columns.id_miembro == id_miembro,
            miembros_comisiones.columns.fecha_incorporacion.between(
                fecha_inicio, fecha_fin
            ),
        )
    ).fetchall()

    # Si ese miembro no pertenece a ninguna comisión no generamos certificado
    if not comisiones:
        flash(
            f"El miembro {miembro.nombre} {miembro.apellidos} no ha pertenecido a ninguna comisión en el rango de fechas indicado.",
            "error",
        )
        return redirect(url_for("informe.generate_informe"))

    # Obtenemos el nombre de esas comisiones a partir de la ID
    # Primero obtenemos todas las ID de las comisiones de la lista anterior
    id_comisiones = [c[0] for c in comisiones]
    # Ahora recuperamos los nombres
    nombres_comisiones = (
        db.session.query(Comision.id, Comision.nombre)
        .filter(Comision.id.in_(id_comisiones))
        .all()
    )
    # Creamos una lista de tuplas con las comisiones, fechas y nombres utilizando map() para agregar el nombre de la comisión
    comisiones_con_nombre = list(
        map(
            lambda c: (
                c[0],
                c[1],
                c[2],
                c[3],
                next((n[1] for n in nombres_comisiones if n[0] == c[0]), ""),
            ),
            comisiones,
        )
    )

    # Ordenamos la lista de manera ascendente en función de la fecha de inicio
    comisiones_con_nombre = sorted(comisiones_con_nombre, key=lambda tupla: tupla[1])

    # Borramos la comisión 'Junta de Escuela' si existe ya que no pertenece a este informe
    for tupla in comisiones_con_nombre:
        if tupla[3] == "Junta de Escuela":
            comisiones_con_nombre.remove(tupla)

    comisiones = comisiones_con_nombre

    fecha_actual = datetime.datetime.now().date()

    # --- Contenido documento word --- #
    document = Document()

    # Modificamos metadatos
    core_properties = document.core_properties

    core_properties.author = "Secretaría EPSC"
    core_properties.title = "Certificado de pertenencia a comisiones"
    core_properties.comments = ""

    # Establecemos el estilo de fuente predeterminado para el documento
    default_font = document.styles["Normal"].font
    # default_font.name = "Arial"
    default_font.size = Pt(12)

    # --- #
    p = document.add_paragraph()
    run = p.add_run()
    run.add_picture(
        os.path.join(BASE_DIR, "app", "static", "images", "logotipo-EPSC.png"),
        width=Inches(3.25),
    )

    # Ajustamos la posición de la imagen
    p.paragraph_format.space_before = Pt(0)

    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # --- Logo EPSC como encabezado --- #
    # Escogemos la sección más arriba de la página
    # section = document.sections[0]

    # Seleccionamos el header
    # header = section.header

    # Seleccionamos el párrafo que hay en la sección del header
    # header_para = header.paragraphs[0]

    # Añadimos el logo
    # run = header_para.add_run()
    # run.add_picture(
    #    os.path.join(BASE_DIR, "app", "static", "images", "logotipo-EPSC.png"),
    #    width=Inches(3.25),
    # )

    # Ajustar el espaciado antes del párrafo para reducir la distancia desde la parte superior
    # header_para.paragraph_format.space_before = Pt(0)

    # header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # --- Logo EPSC como encabezado --- #

    # --- #

    p = document.add_paragraph()
    p.add_run(
        f"{secretario.upper()} DE  LA  ESCUELA POLITÉCNICA SUPERIOR DE CÓRDOBA DE LA UNIVERSIDAD DE CÓRDOBA"
    ).bold = True

    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p.add_run().add_break()

    p = document.add_paragraph()
    p.add_run("CERTIFICA:").bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = document.add_paragraph("Que ")
    p.add_run(
        f"{tratamiento.upper()} {miembro.nombre.upper()} {miembro.apellidos.upper()}"
    ).bold = True
    p.add_run(
        f", con DNI {miembro.dni}, ha sido miembro de las siguientes Comisiones de la Escuela Politécnica Superior de Córdoba y durante los periodos:"
    )

    p.add_run().add_break()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # Indentamos primera línea del párrafo
    paragraph_format = p.paragraph_format
    paragraph_format.first_line_indent = Inches(0.5)

    # Lista de comisiones
    for comision in comisiones:
        if not comision[2]:
            if comision[3]:
                p = document.add_paragraph(
                    f"{comision[4].upper()} desde el {comision[1].strftime('%d/%m/%Y')} hasta la actualidad (en calidad de {comision[3].upper()} de la Comisión).",
                    style="List Bullet",
                )
            else:
                p = document.add_paragraph(
                    f"{comision[4].upper()} desde el {comision[1].strftime('%d/%m/%Y')} hasta la actualidad.",
                    style="List Bullet",
                )
        else:
            if comision[3]:
                p = document.add_paragraph(
                    f"{comision[4].upper()} desde el {comision[1].strftime('%d/%m/%Y')} hasta el {comision[2].strftime('%d/%m/%Y')} (en calidad de {comision[3].upper()} de la Comisión).",
                    style="List Bullet",
                )
            else:
                p = document.add_paragraph(
                    f"{comision[4].upper()} desde el {comision[1].strftime('%d/%m/%Y')} hasta el {comision[2].strftime('%d/%m/%Y')}.",
                    style="List Bullet",
                )

        p.add_run().add_break()

        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # La lista de comisiones tienen un punto menos de tamaño
        p.style.font.size = Pt(11)

    document.add_page_break()

    p = document.add_paragraph()
    p.add_run(
        f"Y, para que conste y surta los efectos oportunos, firmo el presente certificado."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # Indentamos primera línea del párrafo
    paragraph_format = p.paragraph_format
    paragraph_format.first_line_indent = Inches(0.5)

    p = document.add_paragraph()

    # Ajustamos el margen superior del párrafo
    paragraph_format = p.paragraph_format
    paragraph_format.space_before = Inches(
        5
    )  # Ajustamos la altura de la página para que añadamos la imagen y se encuentre al final
    paragraph_format.left_indent = Inches(
        -1
    )  # Ajustamos el margen izquierda para que quede centrada

    run = p.add_run()
    run.add_picture(
        os.path.join(BASE_DIR, "app", "static", "images", "footer-uco.png"),
        width=Inches(8),
    )

    # --- Fin contenido --- #

    # Guardamos el documento en un objeto BytesIO
    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    # Creamos la respuesta del archivo adjunto
    response = make_response(output.getvalue())
    apellidos = "_".join(miembro.apellidos.split(" "))
    filename = f"certificado_comisiones_{miembro.nombre}_{apellidos}.docx"
    response.headers["Content-Disposition"] = f"attachment; filename={filename}"
    response.headers[
        "Content-Type"
    ] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return response


@informe_bp.route(
    "/informe/escuela/docx/<secretario>/<fecha_inicio>/<fecha_fin>/<tratamiento>/<int:id_miembro>"
)
@login_required
def informe_escuela_docx(secretario, fecha_inicio, fecha_fin, tratamiento, id_miembro):
    """
    Ruta que genera un informe de pertenencia a la junta de escuela de un miembro en formato '.docx'
    """
    miembro = Miembro.get_by_id(id_miembro)

    # Aquí sólo nos importa la comisión "Junta de Escuela"
    # Recuperamos el ID de esta comisión
    id_comision = Comision.query.filter_by(nombre="Junta de Escuela").first().id

    if not id_comision:
        flash(
            "Parece que la comisión 'Junta de Escuela' no se encuentra registrada en el sistema.",
            "error",
        )
        return redirect(url_for("informe.generate_informe"))
    # Si este miembro está en esa comsión recuperamos sus fechas de incorporación y de baja
    # Ahora la fecha de incorporación no es necesario que esté entre los rangos que nos indica
    # Ahora nos es irrelevante las fechas que nos indique
    comision = db.session.execute(
        select(
            miembros_comisiones.columns.fecha_incorporacion,
            miembros_comisiones.columns.fecha_baja,
        ).where(
            miembros_comisiones.columns.id_miembro == id_miembro,
            miembros_comisiones.columns.id_comision == id_comision,
        )
    ).first()

    if not comision:
        flash(
            f"Parece que el miembro {miembro.nombre} {miembro.apellidos} no ha pertenecido a la Junta de la Escuela en el rango de fechas indicado.",
            "error",
        )
        return redirect(url_for("informe.generate_informe"))

    fecha_actual = datetime.datetime.now().date()

    # --- Contenido documento word --- #
    document = Document()

    # Modificamos metadatos
    core_properties = document.core_properties

    core_properties.author = "Secretaría EPSC"
    core_properties.title = "Certificado de pertenencia a la Junta de Escuela"
    core_properties.comments = ""

    # Establecemos el estilo de fuente predeterminado para el documento
    default_font = document.styles["Normal"].font
    # default_font.name = "Arial"
    default_font.size = Pt(12)

    # -- LOGO -- #
    # p = document.add_paragraph()
    # run = p.add_run()
    # run.add_picture(
    #    os.path.join(BASE_DIR, "app", "static", "images", "logotipo-EPSC.png"),
    #    width=Inches(3.25),
    # )
    # p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Choosing the top most section of the page
    section = document.sections[0]

    # Selecting the header
    header = section.header

    # Selecting the paragraph already present in
    # the header section
    header_para = header.paragraphs[0]

    # Adding the right zoned header
    # header_para.text = "\t\tThis is Right Zoned Header..."
    run = header_para.add_run()
    run.add_picture(
        os.path.join(BASE_DIR, "app", "static", "images", "logotipo-EPSC.png"),
        width=Inches(3.25),
    )

    # Ajustar el espaciado antes del párrafo para reducir la distancia desde la parte superior
    header_para.paragraph_format.space_before = Pt(0)

    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # -- LOGO -- #

    p = document.add_paragraph()
    p.add_run(
        f"{secretario.upper()} DE  LA  ESCUELA POLITÉCNICA SUPERIOR DE CÓRDOBA DE LA UNIVERSIDAD DE CÓRDOBA"
    ).bold = True

    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p.add_run().add_break()

    p = document.add_paragraph()
    p.add_run("CERTIFICA:").bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p = document.add_paragraph("Que ")
    p.add_run(
        f"{tratamiento.upper()} {miembro.nombre.upper()} {miembro.apellidos.upper()}"
    ).bold = True
    if not comision[1]:
        p.add_run(
            f", con DNI {miembro.dni}, es miembro de la Junta de Escuela de la Escuela Politécnica Superior de Córdoba desde {comision[0].strftime('%d/%m/%Y')} hasta la actualidad."
        )
    else:
        p.add_run(
            f", con DNI {miembro.dni}, es miembro de la Junta de Escuela de la Escuela Politécnica Superior de Córdoba desde {comision[0].strftime('%d/%m/%Y')} hasta el {comision[1].strftime('%d/%m/%Y')}."
        )

    p.add_run().add_break()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # Indentamos primera línea del párrafo
    paragraph_format = p.paragraph_format
    paragraph_format.first_line_indent = Inches(0.5)

    p = document.add_paragraph()
    p.add_run(
        f"Y, para que conste y surta los efectos oportunos, firmo el presente certificado."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # Indentamos primera línea del párrafo
    paragraph_format = p.paragraph_format
    paragraph_format.first_line_indent = Inches(0.5)

    p = document.add_paragraph()

    # Ajustamos el margen superior del párrafo
    paragraph_format = p.paragraph_format
    paragraph_format.space_before = Inches(
        2
    )  # Ajustamos la altura de la página para que añadamos la imagen y se encuentre al final
    paragraph_format.left_indent = Inches(
        -1
    )  # Ajustamos el margen izquierda para que quede centrada

    run = p.add_run()
    run.add_picture(
        os.path.join(BASE_DIR, "app", "static", "images", "footer-uco.png"),
        width=Inches(8),
    )

    # --- Fin contenido --- #

    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    # Creamos la respuesta del archivo adjunto
    response = make_response(output.getvalue())
    apellidos = "_".join(miembro.apellidos.split(" "))
    filename = f"certificado_junta_escuela_{miembro.nombre}_{apellidos}.docx"
    response.headers["Content-Disposition"] = f"attachment; filename={filename}"
    response.headers[
        "Content-Type"
    ] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return response


@informe_bp.route("/export/csv")
@login_required
def export_csv():
    """
    Ruta para exportar datos a fichero CSV
    """

    # Obtenemos los datos que queremos exportar al CSV
    result = (
        db.session.query(
            Comision.nombre,
            Miembro.apellidos,
            Miembro.nombre,
            Miembro.dni,
            Miembro.tipo,
            miembros_comisiones.c.fecha_incorporacion,
            miembros_comisiones.c.fecha_baja,
            miembros_comisiones.c.motivo_baja,
            miembros_comisiones.c.cargo,
            Comision.comentarios,
        )
        .join(Miembro, miembros_comisiones.c.id_miembro == Miembro.id)
        .join(Comision, miembros_comisiones.c.id_comision == Comision.id)
        .order_by(Comision.nombre)
    ).all()

    # Exportamos ese resultado a un CSV
    output = io.StringIO()
    writer = csv.writer(output)

    # Insertamos las columnas
    header = [
        "Comisión",
        "Apellidos",
        "Nombre",
        "DNI",
        "Tipo",
        "Fecha de incorporación",
        "Fecha de baja",
        "Motivo de baja",
        "Cargo en la comisión",
        "Comentarios",
    ]
    writer.writerow(header)

    # Insertamos las filas
    for row in result:
        writer.writerow(row)

    # Reiniciamos posición del puntero
    output.seek(0)

    content = output.getvalue()

    # Codificamos en utf-8-sig (Incluye firma BOM para ayudar a programas externos a reconocer UTF-8)
    content = content.encode("utf-8-sig")

    return Response(
        content,
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment;filename=comisiones.csv"},
    )
